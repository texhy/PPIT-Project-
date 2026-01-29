import os
import cv2
from pathlib import Path
from diagrams import extract_diagrams, parse_diagram_bounds, crop_diagrams
from text import extract_text_from_image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


def process_images_to_docx(image_paths, output_docx_path="output.docx"):
    temp_dir = Path("temp_conversion")
    temp_dir.mkdir(exist_ok=True)
    all_content = []
    
    for idx, image_path in enumerate(image_paths, start=1):
        with open(image_path, "rb") as f:
            image_bytes = f.read()
        
        image_filename = Path(image_path).name
        extracted_text = extract_text_from_image(image_bytes, image_filename)
        diagram_result = extract_diagrams(image_path)
        diagram_bounds = parse_diagram_bounds(diagram_result)
        cropped_diagram_paths = []
        if diagram_bounds:
            diagram_output_dir = temp_dir / f"page_{idx}_diagrams"
            cropped_diagram_paths = crop_diagrams(
                image_path, 
                diagram_bounds, 
                output_dir=str(diagram_output_dir)
            )
        all_content.append({
            "page_number": idx,
            "image_name": image_filename,
            "text": extracted_text,
            "diagrams": cropped_diagram_paths
        })
    
    create_docx(all_content, output_docx_path)
    return output_docx_path


def create_docx(content_data, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    for idx, page in enumerate(content_data):
        heading = doc.add_heading(f"Page {page['page_number']}: {page['image_name']}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        if page['text'] and page['text'].strip():
            text_heading = doc.add_heading('Extracted Text', level=2)
            text_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in text_heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            text_lines = page['text'].split('\n')
            for line in text_lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
            
            doc.add_paragraph()
        
        if page['diagrams'] and len(page['diagrams']) > 0:
            diagram_heading = doc.add_heading('Diagrams', level=2)
            diagram_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in diagram_heading.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            for j, diagram_path in enumerate(page['diagrams'], start=1):
                label = doc.add_paragraph(f'Diagram {j}')
                label.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in label.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
                
                try:
                    img = Image.open(diagram_path)
                    width, height = img.size
                    max_width = 6.0
                    max_height = 4.0
                    
                    aspect_ratio = width / height
                    
                    if width > height:
                        if width > max_width * 96:  
                            display_width = max_width
                            display_height = display_width / aspect_ratio
                        else:
                            display_width = width / 96
                            display_height = height / 96
                    else:
 
                        if height > max_height * 96:
                            display_height = max_height
                            display_width = display_height * aspect_ratio
                        else:
                            display_width = width / 96
                            display_height = height / 96

                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(diagram_path, width=Inches(display_width))
                    
                except Exception as e:
                    error_para = doc.add_paragraph(f'[Error loading diagram: {str(e)}]')
                    error_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in error_para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(255, 0, 0)

                doc.add_paragraph()
        
        if idx < len(content_data) - 1:
            doc.add_page_break()
    
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    sample_images = [
        rf"D:\uni\sem8\ppit\2.jpeg"
    ]
    
    output_file = "handwritten_notes.docx"
    process_images_to_docx(sample_images, output_file)
